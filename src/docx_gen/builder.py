"""
builder.py
Generates a Tamil Word (.docx) document from translated questions.
Produces a clean, print-ready output for coaching institutes.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from loguru import logger

# Tamil Unicode fonts that work well
TAMIL_FONT = "Latha"       # Best for Tamil Nadu institute printing
ENGLISH_FONT = "Calibri"


def set_cell_border(cell):
    """Add borders to table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ["top", "left", "bottom", "right"]:
        tcBorders = OxmlElement(f"w:{side}")
        tcBorders.set(qn("w:val"), "single")
        tcBorders.set(qn("w:sz"), "4")
        tcBorders.set(qn("w:color"), "AAAAAA")
        tcPr.append(tcBorders)


def add_title_page(doc: Document, paper: dict):
    """Add institute header and paper title."""
    # Institute name
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("JEE/NEET தமிழ் மோழிப்பெயர்ப்பு")
    run.font.name = TAMIL_FONT
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x8a)

    # Paper title
    if paper.get("paper_title"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r2 = p2.add_run(paper["paper_title"])
        r2.font.name = TAMIL_FONT
        r2.font.size = Pt(14)
        r2.font.bold = True

    # Subject + Date line
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subject = paper.get("subject", "")
    date_str = datetime.now().strftime("%d-%m-%Y")
    r3 = p3.add_run(f"புல ம்: {subject}    |    தேதி: {date_str}")
    r3.font.name = TAMIL_FONT
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Divider line
    doc.add_paragraph("_" * 80)
    doc.add_paragraph()  # spacer


def add_question(doc: Document, question: dict, q_index: int):
    """Add a single question with options to the document."""

    # ── Question Number + Body ─────────────────────────────────────────
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_before = Pt(8)
    q_para.paragraph_format.space_after = Pt(4)

    # Question number in bold blue
    q_num_run = q_para.add_run(f"{q_index}. ")
    q_num_run.font.name = ENGLISH_FONT
    q_num_run.font.size = Pt(12)
    q_num_run.font.bold = True
    q_num_run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x8a)

    # Question body in Tamil
    q_body_run = q_para.add_run(question.get("body", ""))
    q_body_run.font.name = TAMIL_FONT
    q_body_run.font.size = Pt(12)

    # ── Options A / B / C / D ─────────────────────────────────────────
    options = question.get("options", {})
    if options:
        # Use a 2-column layout for options (A, B on one row; C, D on next)
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.allow_autofit = True

        option_labels = ["A", "B", "C", "D"]
        positions = [(0, 0), (0, 1), (1, 0), (1, 1)]  # (row, col)

        for label, (row, col) in zip(option_labels, positions):
            if label in options:
                cell = table.cell(row, col)
                cell.width = Inches(3)
                para = cell.paragraphs[0]

                # Option label
                label_run = para.add_run(f"({label}) ")
                label_run.font.name = ENGLISH_FONT
                label_run.font.size = Pt(11)
                label_run.font.bold = True

                # Option text
                text_run = para.add_run(options[label])
                text_run.font.name = TAMIL_FONT
                text_run.font.size = Pt(11)

        doc.add_paragraph()  # small gap after table

    # ── Explanation (if present) ──────────────────────────────────
    if question.get("explanation"):
        exp_para = doc.add_paragraph()
        exp_run = exp_para.add_run(f"விளக்கம்: {question['explanation']}")
        exp_run.font.name = TAMIL_FONT
        exp_run.font.size = Pt(10)
        exp_run.font.italic = True
        exp_run.font.color.rgb = RGBColor(0x44, 0x77, 0x44)


def build_docx(paper: dict, output_path: str) -> str:
    """
    Main function: build complete Tamil Word document from paper dict.
    Returns the path to the saved .docx file.
    """
    logger.info(f"Building DOCX: {output_path}")

    doc = Document()

    # Page margins (A4)
    section = doc.sections[0]
    section.page_height = Pt(841)
    section.page_width = Pt(595)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Add title page
    add_title_page(doc, paper)

    # Add all questions
    questions = paper.get("questions", [])
    for i, question in enumerate(questions, 1):
        add_question(doc, question, i)

    # Footer: total questions count
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run(
        f"மொத்த வினாக்கள்: {len(questions)}    |"
        f"    மொத்த மதிப்பெண்கள்: {len(questions) * 4}"
    )
    footer_run.font.name = TAMIL_FONT
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    logger.success(f"DOCX saved: {output_path}")
    return output_path
